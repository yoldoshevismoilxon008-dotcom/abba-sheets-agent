"""kb.py — Bilim bazasi (Knowledge Base) yadrosi (B1).

SQLite FTS5 + Claude re-rank. Embedding yo'q (v1). DB: DATA/knowledge.db.

Ommaviy API:
    init_db()                              — sxema (idempotent, har startda)
    ingest_text(text, *, source, origin, ...) -> dict
    ingest_file(path, *, source, origin, caption="") -> dict   # content_key=True
    archive(uid_or_title) -> dict          — /unut (ok|not_found|ambiguous; ommaviy emas)
    stats() -> dict
    search(query, k=8, *, use_rerank=True, use_expansion=True) -> list[dict]
    context_for(query, budget_chars=12_000, *, use_rerank=True, use_expansion=True) -> str
    list_docs(limit=30, tag=None) -> list[dict]

Qat'iy tamoyil: KB Claude sababli YIQILMAYDI — metadata/query/re-rank
har biri try/except, fallback bor. Ingest hech qachon Claude xatosidan to'xtamaydi.
"""

import os
import re
import json
import sqlite3
import hashlib
import mimetypes
from pathlib import Path
from datetime import datetime

BASE = Path(__file__).resolve().parent
DATA = Path(os.environ.get("DATA_DIR") or (BASE / "data"))
DB_PATH = DATA / "knowledge.db"
PROMPTS = BASE / "prompts"

SCHEMA_VERSION = 1
MAX_CHARS = 3000            # bo'lak maksimal hajmi
OVERLAP = 300              # katta bo'lak bo'linganda ustma-ust
MERGE_MIN = 200            # kichik bo'lak keyingisiga qo'shiladi
META_SNIPPET = 6000       # metadata uchun birinchi N belgi
CTX_BUDGET = 12_000       # context_for standart budjet
MAX_EXTRACT_CHARS = 4_000_000   # bitta hujjatdan olinadigan maks matn (xavfsizlik)
MAX_CHUNKS = 2000         # bitta hujjat maks bo'lak soni (oshsa kesiladi)
CLAUDE_TIMEOUT = 30       # KB Claude chaqiruvlari uchun qattiq timeout (sekund)
# Chat arxivi (B2.2) — botning avvalgi javoblari, PAST vaznli kontekst
CHAT_SCORE_PENALTY = float(os.environ.get("KB_CHAT_PENALTY", "0.6"))  # search: chat ballini demote
CHAT_CTX_MAX = int(os.environ.get("KB_CHAT_CTX_MAX", "2"))            # context_for: ko'pi 2 chat bo'lagi

# O'zbek/tipografik apostrof variantlari → bitta ASCII ' (index va query BIR XIL)
_APOS = {
    "ʻ": "'", "ʼ": "'", "‘": "'", "’": "'",
    "`": "'", "´": "'", "ʹ": "'", "′": "'",
}
_APOS_RE = re.compile("[" + "".join(_APOS) + "]")


def _canon_apos(s):
    """Barcha apostrof/tipografik variantlarni ASCII ' ga keltiradi.
    _normalize (saqlash/indeks) va _fts_query (so'rov) — IKKALASIDA bir xil."""
    return _APOS_RE.sub("'", s or "")


def log(msg):
    print(f"[kb] {msg}", flush=True)


def _now():
    return datetime.now().isoformat(timespec="seconds")


# ======================================================================
# DB — sxema va ulanish
# ======================================================================

_SCHEMA = """
CREATE TABLE IF NOT EXISTS docs (
  id           INTEGER PRIMARY KEY,
  uid          TEXT UNIQUE NOT NULL,
  title        TEXT NOT NULL,
  source       TEXT NOT NULL,
  origin       TEXT NOT NULL,
  mime         TEXT,
  lang         TEXT,
  tags         TEXT,
  summary      TEXT,
  content_hash TEXT NOT NULL,
  vault_path   TEXT,
  bytes        INTEGER,
  n_chunks     INTEGER DEFAULT 0,
  created_at   TEXT NOT NULL,
  updated_at   TEXT NOT NULL,
  archived     INTEGER DEFAULT 0
);

CREATE TABLE IF NOT EXISTS chunks (
  id        INTEGER PRIMARY KEY,
  doc_id    INTEGER NOT NULL REFERENCES docs(id) ON DELETE CASCADE,
  ord       INTEGER NOT NULL,
  heading   TEXT,
  text      TEXT NOT NULL,
  n_chars   INTEGER,
  embedding BLOB
);
CREATE INDEX IF NOT EXISTS idx_chunks_doc ON chunks(doc_id);

CREATE VIRTUAL TABLE IF NOT EXISTS chunks_fts USING fts5(
  text, heading,
  content='chunks', content_rowid='id',
  tokenize='unicode61 remove_diacritics 2'
);

CREATE TRIGGER IF NOT EXISTS chunks_ai AFTER INSERT ON chunks BEGIN
  INSERT INTO chunks_fts(rowid, text, heading) VALUES (new.id, new.text, new.heading);
END;
CREATE TRIGGER IF NOT EXISTS chunks_ad AFTER DELETE ON chunks BEGIN
  INSERT INTO chunks_fts(chunks_fts, rowid, text, heading) VALUES('delete', old.id, old.text, old.heading);
END;
CREATE TRIGGER IF NOT EXISTS chunks_au AFTER UPDATE ON chunks BEGIN
  INSERT INTO chunks_fts(chunks_fts, rowid, text, heading) VALUES('delete', old.id, old.text, old.heading);
  INSERT INTO chunks_fts(rowid, text, heading) VALUES (new.id, new.text, new.heading);
END;

CREATE TABLE IF NOT EXISTS ingest_log (
  id INTEGER PRIMARY KEY, ts TEXT, source TEXT, origin TEXT,
  status TEXT, n_chunks INTEGER, err TEXT
);
"""


def _connect():
    DATA.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(str(DB_PATH))
    conn.execute("PRAGMA foreign_keys=ON")
    conn.execute("PRAGMA journal_mode=WAL")
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    """Sxemani yaratadi (idempotent). Har startda chaqirsa bo'ladi."""
    conn = _connect()
    try:
        conn.executescript(_SCHEMA)
        ver = conn.execute("PRAGMA user_version").fetchone()[0]
        if ver < SCHEMA_VERSION:
            conn.execute(f"PRAGMA user_version = {SCHEMA_VERSION}")
        conn.commit()
    finally:
        conn.close()
    return True


# ======================================================================
# Matn: normalize, hash
# ======================================================================

def _normalize(text):
    """Deterministik normalize — content_hash barqaror bo'lishi uchun.
    Apostrof variantlari ASCII ' ga (qidiruv cross-match ishlashi uchun)."""
    text = _canon_apos((text or "").replace("\r\n", "\n").replace("\r", "\n"))
    out, blank = [], 0
    for ln in text.split("\n"):
        ln = ln.rstrip()
        if ln == "":
            blank += 1
            if blank <= 2:
                out.append(ln)
        else:
            blank = 0
            out.append(ln)
    return "\n".join(out).strip()


def _sha(s):
    return hashlib.sha256(s.encode("utf-8")).hexdigest()


def _uid(source, origin, extra=None):
    """Hujjat identifikatori. extra (odatda content_hash) berilsa — kontent-manzilli:
    bir xil origin (masalan ikkita boshqa 'document.pdf') endi bitta uid olmaydi,
    birinchisining chunk'lari o'chib ketmaydi (FIX: uid kolliziyasi)."""
    key = f"{source}:{origin}"
    if extra:
        key += f":{extra}"
    return _sha(key)[:16]


# ======================================================================
# Chunking — markdown sarlavhalari bo'yicha, katta bo'lsa bo'linadi
# ======================================================================

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*#*\s*$")


def _pick_heading(a, b):
    """Ikki sarlavha yo'lidan mosini tanlaydi (chuqurroq/ota-onasini)."""
    if not a:
        return b
    if not b:
        return a
    if b.startswith(a):
        return b
    if a.startswith(b):
        return a
    return a


def _sections(text):
    """Matnni markdown sarlavhalari bo'yicha (path, body) bo'laklarga ajratadi."""
    sections, buf, stack, path = [], [], [], ""

    def flush():
        body = "\n".join(buf).strip()
        if body:
            sections.append((path, body))
        buf.clear()

    for line in text.split("\n"):
        m = _HEADING_RE.match(line)
        if m:
            flush()
            level = len(m.group(1))
            title = m.group(2).strip()
            while stack and stack[-1][0] >= level:
                stack.pop()
            stack.append((level, title))
            path = " > ".join(t for _, t in stack)
            buf.append(line)   # sarlavha yangi bo'lak boshiga kiradi
        else:
            buf.append(line)
    flush()
    return sections


def _split_large(text, max_chars=MAX_CHARS, overlap=OVERLAP):
    """Katta bo'lakni paragraf chegarasida bo'ladi; jadval satri o'rtasidan kesilmaydi."""
    if len(text) <= max_chars:
        return [text]
    # 1) paragraflarga; juda katta paragraf bo'lsa — satr chegarasida (mid-row emas)
    units = []
    for para in re.split(r"\n\s*\n", text):
        if len(para) <= max_chars:
            units.append(para)
            continue
        cur = ""
        for line in para.split("\n"):
            if cur and len(cur) + len(line) + 1 > max_chars:
                units.append(cur)
                cur = line
            else:
                cur = (cur + "\n" + line) if cur else line
        if cur:
            units.append(cur)
    # 1b) xavfsizlik: hali ham max'dan katta unit (uzun bitta satr, split nuqtasi yo'q)
    #     — belgilar bo'yicha majburiy bo'lamiz (aks holda ulkan chunk chiqadi)
    safe = []
    for u in units:
        if len(u) <= max_chars:
            safe.append(u)
        else:
            for i in range(0, len(u), max_chars):
                safe.append(u[i:i + max_chars])
    units = safe
    # 2) unitlarni max_chars gacha yig'ish, ustma-ust (overlap) bilan
    pieces, cur = [], ""
    for u in units:
        if cur and len(cur) + len(u) + 2 > max_chars:
            pieces.append(cur)
            tail = cur[-overlap:] if overlap else ""
            cur = (tail + "\n\n" + u) if tail else u
        else:
            cur = (cur + "\n\n" + u) if cur else u
    if cur:
        pieces.append(cur)
    return pieces


def chunk_text(text):
    """Normalize qilingan matndan bo'laklar ro'yxati:
       [{ord, heading, text, n_chars}, ...]"""
    text = _normalize(text)
    sections = _sections(text)
    if not sections:
        return []
    # kichik bo'laklarni keyingisiga qo'shish (forward merge)
    merged, pending = [], None
    for path, body in sections:
        if pending:
            body = pending[1] + "\n\n" + body
            path = _pick_heading(pending[0], path)
            pending = None
        if len(body) < MERGE_MIN:
            pending = (path, body)
            continue
        merged.append((path, body))
    if pending:
        if merged:
            pp, pb = merged[-1]
            merged[-1] = (_pick_heading(pp, pending[0]), pb + "\n\n" + pending[1])
        else:
            merged.append(pending)
    # katta bo'laklarni bo'lish
    final = []
    for path, body in merged:
        for piece in _split_large(body):
            final.append((path, piece))
    return [
        {"ord": i, "heading": p, "text": t, "n_chars": len(t)}
        for i, (p, t) in enumerate(final)
    ]


# ======================================================================
# Extractorlar
# ======================================================================

_TEXT_EXTS = {"md", "markdown", "txt", "text", "csv", "tsv", "log"}


def _read_text_file(path):
    return path.read_text(encoding="utf-8", errors="replace")


def _extract_pdf(path):
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    parts, total = [], 0
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        if t.strip():
            parts.append(t.strip())
            total += len(t)
            if total >= MAX_EXTRACT_CHARS:   # ulkan PDF — xotira/timeout xavfsizligi
                break
    return "\n\n".join(parts)


def _extract_docx(path):
    import docx
    d = docx.Document(str(path))
    parts, total = [], 0
    for p in d.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text)
            total += len(p.text)
            if total >= MAX_EXTRACT_CHARS:
                return "\n\n".join(parts)
    for tbl in d.tables:
        for row in tbl.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            if any(cells):
                line = " | ".join(cells)
                parts.append(line)
                total += len(line)
                if total >= MAX_EXTRACT_CHARS:
                    return "\n\n".join(parts)
    return "\n\n".join(parts)


def _extract_html(path):
    from bs4 import BeautifulSoup
    raw = path.read_text(encoding="utf-8", errors="replace")
    try:
        soup = BeautifulSoup(raw, "lxml")
    except Exception:
        soup = BeautifulSoup(raw, "html.parser")   # lxml bo'lmasa — Docker o'zgarmaydi
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text("\n")


def _extract_xlsx(path):
    from openpyxl import load_workbook
    wb = load_workbook(str(path), read_only=True, data_only=True)
    parts, total = [], 0
    try:
        for ws in wb.worksheets:
            parts.append(f"## {ws.title}")
            for row in ws.iter_rows(values_only=True):
                cells = ["" if v is None else str(v) for v in row]
                if any(c.strip() for c in cells):
                    line = ", ".join(cells)
                    parts.append(line)
                    total += len(line)
                    if total >= MAX_EXTRACT_CHARS:   # ulkan xlsx — xavfsizlik
                        return "\n".join(parts)
    finally:
        wb.close()
    return "\n".join(parts)


def _extract_rtf(path):
    raw = path.read_text(encoding="utf-8", errors="replace")
    txt = re.sub(r"\\par[d]?", "\n", raw)
    txt = re.sub(r"\\[a-zA-Z]+-?\d* ?", "", txt)
    return txt.replace("{", "").replace("}", "").strip()


def extract_file(path):
    """Kengaytmaga qarab matn ajratadi. Noma'lum — matn sifatida o'qiladi.
    Yakuniy matn MAX_EXTRACT_CHARS bilan cheklanadi (chunk soni ceiling'i ingest_text'da)."""
    path = Path(path)
    ext = path.suffix.lower().lstrip(".")
    if ext == "pdf":
        text = _extract_pdf(path)
    elif ext == "docx":
        text = _extract_docx(path)
    elif ext in ("html", "htm"):
        text = _extract_html(path)
    elif ext == "xlsx":
        text = _extract_xlsx(path)
    elif ext == "rtf":
        text = _extract_rtf(path)
    else:   # md/txt/csv/tsv/log/noma'lum — matn sifatida
        text = _read_text_file(path)
    return text[:MAX_EXTRACT_CHARS] if text and len(text) > MAX_EXTRACT_CHARS else text


# ======================================================================
# Claude yordamchilari (metadata, query expansion, re-rank) — hammasi fallback bilan
# ======================================================================

def _ask_claude(prompt, effort="low"):
    """analyze.run_claude ustidan yupqa qatlam (test uchun monkeypatch qulay).
    QATTIQ timeout=CLAUDE_TIMEOUT (30s) — osilgan chaqiruv botni bloklamasin;
    timeout/xato → chaqiruvchi fallback qiladi (meta→fayl nomi, query→tokenizatsiya,
    rerank→BM25 tartibi)."""
    import analyze
    return analyze.run_claude(prompt, effort=effort, timeout=CLAUDE_TIMEOUT)


def _load_prompt(name):
    return (PROMPTS / name).read_text(encoding="utf-8")


def _parse_json_obj(raw):
    m = re.search(r"\{.*\}", raw or "", re.DOTALL)
    if not m:
        raise ValueError("JSON obyekt topilmadi")
    return json.loads(m.group(0))


def _parse_json_arr(raw):
    m = re.search(r"\[.*\]", raw or "", re.DOTALL)
    if not m:
        raise ValueError("JSON massiv topilmadi")
    return json.loads(m.group(0))


def _fallback_title(origin):
    stem = Path(origin).stem if origin else ""
    stem = re.sub(r"[-_]+", " ", stem).strip()
    return (stem[:120] or "Hujjat")


def _first_heading(text):
    """Matndagi birinchi markdown sarlavhasi (# ...) matni; bo'lmasa ''."""
    for line in (text or "").split("\n"):
        m = _HEADING_RE.match(line)
        if m:
            return m.group(2).strip()
    return ""


def _light_meta(text, origin, title_hint=None):
    """Claude'SIZ metadata — kichik fayllar uchun (narx tejash, meta_min_chars).
    Sarlavha: title_hint → birinchi markdown sarlavha → fayl nomidan fallback."""
    title = (title_hint or "").strip() or _first_heading(text) or _fallback_title(origin)
    return {"title": title[:200], "lang": "uz", "tags": [], "summary": text.strip()[:300]}


def _chat_meta(text, origin):
    """Chat arxivi metadata — Claude UMUMAN chaqirilmaydi (B2.2 fix#1). Sabab: joriy kun
    fayli kun davomida O'SADI → har 10 daqiqada content_hash o'zgaradi → 1500 belgi
    chegarasidan oshgani uchun Claude metadata kuniga ~144 marta BEKORGA chaqirilardi.
    Doim barqaror fallback sarlavha «Suhbat arxivi — YYYY-MM-DD»."""
    m = re.search(r"\d{4}-\d{2}-\d{2}", Path(origin).name if origin else "")
    day = m.group(0) if m else (Path(origin).stem if origin else "")
    return {"title": f"Suhbat arxivi — {day}".strip()[:200], "lang": "uz",
            "tags": ["suhbat"], "summary": text.strip()[:300]}


def _extract_metadata(text, origin, title_hint=None):
    """Claude bilan {title, lang, tags, summary}. Yiqilsa — fallback."""
    try:
        prompt = (
            _load_prompt("kb-meta.md")
            .replace("{{FILENAME}}", origin or "")
            .replace("{{TEXT}}", text[:META_SNIPPET])
        )
        data = _parse_json_obj(_ask_claude(prompt))
        title = title_hint or (data.get("title") or "").strip() or _fallback_title(origin)
        lang = (data.get("lang") or "").strip().lower() or "uz"
        tags = [str(t).strip().lower() for t in (data.get("tags") or []) if str(t).strip()][:8]
        summary = (data.get("summary") or "").strip()
        return {"title": title[:200], "lang": lang, "tags": tags, "summary": summary}
    except Exception as e:
        log(f"metadata fallback ({origin}): {e}")
        return {
            "title": (title_hint or _fallback_title(origin))[:200],
            "lang": "uz", "tags": [], "summary": text.strip()[:300],
        }


# ======================================================================
# Yozish — ingest
# ======================================================================

def _log_ingest(conn, source, origin, status, n_chunks, err=None):
    try:
        conn.execute(
            "INSERT INTO ingest_log (ts, source, origin, status, n_chunks, err) VALUES (?,?,?,?,?,?)",
            (_now(), source, origin, status, n_chunks, err),
        )
    except Exception:
        pass


def _safe_log_ingest(source, origin, status, n_chunks, err):
    try:
        conn = _connect()
        _log_ingest(conn, source, origin, status, n_chunks, err)
        conn.commit()
        conn.close()
    except Exception:
        pass


def ingest_text(text, *, source, origin, title=None, mime="text/plain",
                tags=None, vault_path=None, content_key=False, meta_min_chars=0):
    """Matnni normalize → chunk → metadata (Claude) → upsert.
    content_key=True → uid kontent-manzilli (bir xil origin'li boshqa kontent
    ustiga yozilmaydi; ingest_file/eslatma shu rejimda).
    meta_min_chars>0 → shu belgidan KICHIK matnga Claude metadata chaqirilmaydi,
    light meta (fayl nomi + birinchi sarlavha) ishlatiladi (vault narx-nazorati).
    Qaytadi: {uid, title, n_chunks, status: new|updated|unchanged, tags, summary, warn}"""
    init_db()
    norm = _normalize(text)
    if not norm:
        raise ValueError("bo'sh matn — ingest qilinmadi")
    warn_parts = []
    if len(norm) > MAX_EXTRACT_CHARS:
        norm = norm[:MAX_EXTRACT_CHARS]
        warn_parts.append(f"matn {MAX_EXTRACT_CHARS} belgida kesildi")
    chash = _sha(norm)
    uid = _uid(source, origin, chash if content_key else None)
    now = _now()

    conn = _connect()
    try:
        row = conn.execute(
            "SELECT id, content_hash, title, tags, summary, n_chunks FROM docs WHERE uid=?",
            (uid,),
        ).fetchone()

        if row and row["content_hash"] == chash and row["n_chunks"]:
            _log_ingest(conn, source, origin, "unchanged", row["n_chunks"])
            conn.commit()
            return {
                "uid": uid, "title": row["title"], "n_chunks": row["n_chunks"],
                "status": "unchanged", "tags": json.loads(row["tags"] or "[]"),
                "summary": row["summary"] or "", "warn": "",
            }

        if source == "chat":
            meta = _chat_meta(norm, origin)          # chat — Claude YO'Q (fix#1)
        elif meta_min_chars and len(norm) < meta_min_chars:
            meta = _light_meta(norm, origin, title_hint=title)
        else:
            meta = _extract_metadata(norm, origin, title_hint=title)
        if tags:
            extra = [str(t).strip().lower() for t in tags if str(t).strip()]
            meta["tags"] = list(dict.fromkeys([*extra, *meta["tags"]]))[:10]

        chunks = chunk_text(norm)
        if not chunks:
            raise ValueError("chunk chiqmadi (bo'sh matn)")
        if len(chunks) > MAX_CHUNKS:
            warn_parts.append(f"{len(chunks)} bo'lakdan {MAX_CHUNKS} tasi saqlandi")
            chunks = chunks[:MAX_CHUNKS]

        tags_json = json.dumps(meta["tags"], ensure_ascii=False)
        nbytes = len(norm.encode("utf-8"))

        if row:
            doc_id = row["id"]
            conn.execute("DELETE FROM chunks WHERE doc_id=?", (doc_id,))
            conn.execute(
                """UPDATE docs SET title=?, source=?, origin=?, mime=?, lang=?, tags=?,
                   summary=?, content_hash=?, vault_path=?, bytes=?, n_chunks=?,
                   updated_at=?, archived=0 WHERE id=?""",
                (meta["title"], source, origin, mime, meta["lang"], tags_json,
                 meta["summary"], chash, vault_path, nbytes, len(chunks), now, doc_id),
            )
            status = "updated"
        else:
            cur = conn.execute(
                """INSERT INTO docs (uid,title,source,origin,mime,lang,tags,summary,
                   content_hash,vault_path,bytes,n_chunks,created_at,updated_at,archived)
                   VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,0)""",
                (uid, meta["title"], source, origin, mime, meta["lang"], tags_json,
                 meta["summary"], chash, vault_path, nbytes, len(chunks), now, now),
            )
            doc_id = cur.lastrowid
            status = "new"

        conn.executemany(
            "INSERT INTO chunks (doc_id, ord, heading, text, n_chars) VALUES (?,?,?,?,?)",
            [(doc_id, c["ord"], c["heading"], c["text"], c["n_chars"]) for c in chunks],
        )
        _log_ingest(conn, source, origin, status, len(chunks))
        conn.commit()
        return {
            "uid": uid, "title": meta["title"], "n_chunks": len(chunks),
            "status": status, "tags": meta["tags"], "summary": meta["summary"],
            "warn": "; ".join(warn_parts),
        }
    except Exception as e:
        conn.rollback()
        _log_ingest(conn, source, origin, "error", 0, str(e)[:400])
        try:
            conn.commit()
        except Exception:
            pass
        raise
    finally:
        conn.close()


def ingest_file(path, *, source, origin, caption="", vault_path=None,
                title=None, meta_min_chars=0):
    """Kengaytmaga qarab extractor → ingest_text (content_key=True —
    bir xil nomli boshqa fayl birinchisining chunk'larini o'chirmasin).
    vault_path/title/meta_min_chars → ingest_text ga uzatiladi (vault_sync uchun)."""
    path = Path(path)
    text = extract_file(path)
    if not text or not text.strip():
        raise ValueError("fayldan matn ajratib bo'lmadi (bo'sh yoki skan hujjat)")
    if caption and caption.strip():
        text = f"[Izoh: {caption.strip()}]\n\n{text}"
    mime = mimetypes.guess_type(str(path))[0] or "application/octet-stream"
    return ingest_text(text, source=source, origin=origin, mime=mime,
                       content_key=True, vault_path=vault_path, title=title,
                       meta_min_chars=meta_min_chars)


def forget_origin(source, origin):
    """Aynan shu source+origin hujjat(lar)ini arxivlaydi (vault'da fayl o'chganda).
    Scoped — HECH QACHON ommaviy emas. Qaytadi: arxivlangan hujjat soni."""
    init_db()
    conn = _connect()
    try:
        cur = conn.execute(
            "UPDATE docs SET archived=1, updated_at=? "
            "WHERE source=? AND origin=? AND archived=0",
            (_now(), source, origin),
        )
        conn.commit()
        n = cur.rowcount
    finally:
        conn.close()
    if n:
        _safe_log_ingest(source, origin, "archived", 0, None)
    return n


def _like_escape(s):
    """LIKE meta-belgilarini (\\ % _) neytrallaydi — '%' bilan ommaviy arxivlashning oldi."""
    return s.replace("\\", "\\\\").replace("%", "\\%").replace("_", "\\_")


def archive(uid_or_title):
    """/unut — archived=1 (o'chirmaydi). HECH QACHON ommaviy arxivlamaydi.
    Qaytadi dict:
      {status:"ok", n:1, id, title}                — aniq 1 ta arxivlandi
      {status:"not_found"}                          — mos yo'q
      {status:"ambiguous", candidates:[{id,title}]} — 2+ mos, aniq id kerak"""
    init_db()
    key = str(uid_or_title).strip()
    if not key:
        return {"status": "not_found"}
    conn = _connect()
    try:
        # 1) uid aniq → 2) id aniq (faqat shu ikkisi to'g'ridan-to'g'ri arxivlaydi)
        row = conn.execute(
            "SELECT id, title FROM docs WHERE uid=? AND archived=0", (key,)
        ).fetchone()
        if not row and key.isdigit():
            row = conn.execute(
                "SELECT id, title FROM docs WHERE id=? AND archived=0", (int(key),)
            ).fetchone()
        if row:
            conn.execute(
                "UPDATE docs SET archived=1, updated_at=? WHERE id=?", (_now(), row["id"])
            )
            conn.commit()
            return {"status": "ok", "n": 1, "id": row["id"], "title": row["title"]}
        # 3) title bo'yicha — LIKE escaped, AYNAN 1 ta bo'lsagina arxivlanadi
        rows = conn.execute(
            "SELECT id, title FROM docs WHERE archived=0 AND title LIKE ? ESCAPE '\\' "
            "ORDER BY id LIMIT 11",
            (f"%{_like_escape(key)}%",),
        ).fetchall()
        if not rows:
            return {"status": "not_found"}
        if len(rows) == 1:
            conn.execute(
                "UPDATE docs SET archived=1, updated_at=? WHERE id=?", (_now(), rows[0]["id"])
            )
            conn.commit()
            return {"status": "ok", "n": 1, "id": rows[0]["id"], "title": rows[0]["title"]}
        return {
            "status": "ambiguous",
            "candidates": [{"id": r["id"], "title": r["title"]} for r in rows[:10]],
        }
    finally:
        conn.close()


# ======================================================================
# O'qish — qidiruv
# ======================================================================

def _fallback_keywords(query):
    toks = re.findall(r"\w+", query or "", re.UNICODE)
    seen = []
    for t in toks:
        tl = t.lower()
        if len(tl) >= 3 and tl not in seen:
            seen.append(tl)
    return seen[:8] or [t.lower() for t in toks[:3]]


def _expand_query(query):
    """Claude query expansion (fast). Yiqilsa — savoldan 3+ harfli so'zlar."""
    try:
        prompt = _load_prompt("kb-query.md").replace("{{QUERY}}", query)
        raw = _ask_claude(prompt)
        try:
            arr = _parse_json_arr(raw)
            kws = [str(x).strip() for x in arr if str(x).strip()]
        except Exception:
            kws = [p.strip(" -*•\t") for p in re.split(r"[\n,]+", raw) if p.strip()]
        kws = [k for k in kws if k][:8]
        if kws:
            return kws
    except Exception as e:
        log(f"query expansion fallback: {e}")
    return _fallback_keywords(query)


def _fts_query(keywords):
    """Kalit so'zlar → FTS5 prefix MATCH (har token `token*`, OR bilan).
    Apostrof _normalize bilan BIR XIL kanonikalizatsiya qilinadi (cross-match)."""
    tokens = []
    for kw in keywords:
        for tok in re.findall(r"\w+", _canon_apos(kw), re.UNICODE):
            tl = tok.lower()
            if len(tl) >= 2 and tl not in tokens:
                tokens.append(tl)
    return " OR ".join(f"{t}*" for t in tokens) if tokens else ""


def _rerank(query, cands, k):
    """Claude re-rank — eng mos k tasining indeksi. Yiqilsa → None (BM25 saqlanadi)."""
    try:
        lines = []
        for i, c in enumerate(cands):
            snippet = " ".join((c["text"] or "").split())[:200]
            lines.append(f"[{i}] {c['title']} | {c['heading'] or ''} | {snippet}")
        prompt = (
            _load_prompt("kb-rerank.md")
            .replace("{{QUERY}}", query)
            .replace("{{K}}", str(k))
            .replace("{{CANDIDATES}}", "\n".join(lines))
        )
        arr = _parse_json_arr(_ask_claude(prompt))
        order = []
        for x in arr:
            s = str(x).strip().lstrip("-")
            if s.isdigit():
                order.append(int(x))
        order = list(dict.fromkeys(order))   # dedupe — Claude [0,0,0,1] qaytarsa takror bo'lmasin
        return order[:k] or None
    except Exception as e:
        log(f"re-rank fallback: {e}")
        return None


def search(query, k=8, *, use_rerank=True, use_expansion=True):
    """FTS5 + (ixtiyoriy) Claude re-rank.
    use_expansion=False → Claude query-expansion o'chiq (oddiy tokenizatsiya) —
    fast rejim savollariga latency qo'shmaslik uchun.
    [{chunk_id, doc_id, ord, heading, text, uid, title, origin, source, vault_path, score}]"""
    init_db()
    # Bo'sh KB — Claude'ni ham, so'rovni ham qurmaymiz (fast savollar tez qolsin)
    conn0 = _connect()
    try:
        has_any = conn0.execute("SELECT 1 FROM chunks LIMIT 1").fetchone()
    finally:
        conn0.close()
    if not has_any:
        return []
    keywords = _expand_query(query) if use_expansion else _fallback_keywords(query)
    match = _fts_query(keywords)
    if not match:
        return []
    conn = _connect()
    try:
        rows = conn.execute(
            """SELECT c.id AS chunk_id, c.doc_id, c.ord, c.heading, c.text,
                      d.uid, d.title, d.origin, d.source, d.vault_path,
                      bm25(chunks_fts, 1.0, 0.5)
                        * (CASE WHEN d.source='chat' THEN ? ELSE 1.0 END) AS score
               FROM chunks_fts
               JOIN chunks c ON c.id = chunks_fts.rowid
               JOIN docs d ON d.id = c.doc_id
               WHERE chunks_fts MATCH ? AND d.archived = 0
               ORDER BY score
               LIMIT 30""",
            (CHAT_SCORE_PENALTY, match),
        ).fetchall()
    except sqlite3.OperationalError as e:
        log(f"FTS MATCH xato ({match!r}): {e}")
        return []
    finally:
        conn.close()

    cands = [dict(r) for r in rows]
    if not cands:
        return []
    if use_rerank and len(cands) > k:
        order = _rerank(query, cands, k)
        if order:
            picked = [cands[i] for i in order if 0 <= i < len(cands)]
            return (picked or cands)[:k]
    return cands[:k]


def context_for(query, budget_chars=CTX_BUDGET, *, use_rerank=True, use_expansion=True):
    """Q&A prompt'ga quyiladigan tayyor blok. Bo'sh bo'lsa "" qaytadi.
    use_expansion=False (fast rejim) → Claude query-expansion o'chiq."""
    try:
        results = search(query, k=8, use_rerank=use_rerank, use_expansion=use_expansion)
    except Exception as e:
        log(f"context_for xato: {e}")
        return ""
    if not results:
        return ""
    header = "[BILIM BAZASI — kontekst/qoida, jonli sheet raqami EMAS]\n"
    # B2.2: chat (bot avvalgi javoblari) — vault/doc bo'laklaridan KEYIN, ko'pi CHAT_CTX_MAX,
    # va faqat joy qolsa. Alohida «tasdiqlanmagan» sarlavha ostida (fakt deb keltirilmasin).
    nonchat = [r for r in results if r.get("source") != "chat"]
    chat = [r for r in results if r.get("source") == "chat"][:CHAT_CTX_MAX]
    out, used, added = [header], len(header), 0
    for i, r in enumerate(nonchat):
        title = r.get("title") or "Hujjat"
        head = r.get("heading") or ""
        origin = r.get("origin") or r.get("source") or ""
        vpath = r.get("vault_path") or ""
        # vault manbasi — to'liq nisbiy yo'l (bot javobda xavfsiz <code> blokka aylantiradi)
        manba = vpath if (r.get("source") == "vault" and vpath) else origin
        loc = f"### {title}" + (f" › {head}" if head else "") + f"   (manba: {manba})\n"
        body = (r.get("text") or "").strip()
        block = loc + body + "\n---\n"
        if used + len(block) > budget_chars:
            if i == 0 and added == 0:   # birinchi bo'lak sig'masa — kesib bo'lsa ham qo'shamiz
                avail = max(0, budget_chars - used - len(loc) - 6)
                body = body[:avail].rstrip()
                if body:
                    out.append(loc + body + "\n---\n")
                    added += 1
            break
        out.append(block)
        used += len(block)
        added += 1
    # Chat bo'laklari — faqat vault/doc'dan keyin joy qolsa
    if chat and used < budget_chars:
        chat_hdr = ("\n[O'TGAN SUHBAT — tasdiqlanmagan, fakt sifatida keltirma, "
                    "faqat kontekst uchun]\n")
        started = False
        for r in chat:
            title = r.get("title") or "Suhbat"
            head = r.get("heading") or ""
            loc = f"### {title}" + (f" › {head}" if head else "") + "\n"
            body = (r.get("text") or "").strip()
            extra = (chat_hdr if not started else "") + loc + body + "\n---\n"
            if used + len(extra) > budget_chars:
                break
            out.append(extra)
            used += len(extra)
            added += 1
            started = True
    return "".join(out) if added else ""


# ======================================================================
# Statistika, ro'yxat
# ======================================================================

def stats():
    init_db()
    conn = _connect()
    try:
        d = conn.execute("SELECT COUNT(*) FROM docs WHERE archived=0").fetchone()[0]
        da = conn.execute("SELECT COUNT(*) FROM docs WHERE archived=1").fetchone()[0]
        ch = conn.execute("SELECT COUNT(*) FROM chunks").fetchone()[0]
        by_src = {
            r[0]: r[1]
            for r in conn.execute(
                "SELECT source, COUNT(*) FROM docs WHERE archived=0 GROUP BY source"
            )
        }
        last = conn.execute(
            "SELECT ts, origin, status FROM ingest_log ORDER BY id DESC LIMIT 1"
        ).fetchone()
        last_ingest = {"ts": last[0], "origin": last[1], "status": last[2]} if last else None
    finally:
        conn.close()
    size = DB_PATH.stat().st_size if DB_PATH.exists() else 0
    wal = DB_PATH.with_name(DB_PATH.name + "-wal")
    if wal.exists():
        size += wal.stat().st_size
    return {
        "docs": d, "docs_archived": da, "chunks": ch, "bytes_db": size,
        "by_source": by_src, "last_ingest": last_ingest,
    }


def origins(source):
    """Berilgan source uchun arxivlanmagan hujjat origin'lari (vault reconcile —
    joriy fayllar bilan solishtirib o'chirilganini topish uchun)."""
    init_db()
    conn = _connect()
    try:
        rows = conn.execute(
            "SELECT DISTINCT origin FROM docs WHERE source=? AND archived=0", (source,)
        ).fetchall()
    finally:
        conn.close()
    return [r[0] for r in rows]


def source_counts(source):
    """(docs, chunks) — berilgan source bo'yicha arxivlanmagan hujjat/bo'lak soni."""
    init_db()
    conn = _connect()
    try:
        d = conn.execute(
            "SELECT COUNT(*) FROM docs WHERE source=? AND archived=0", (source,)
        ).fetchone()[0]
        c = conn.execute(
            "SELECT COALESCE(SUM(n_chunks),0) FROM docs WHERE source=? AND archived=0",
            (source,),
        ).fetchone()[0]
    finally:
        conn.close()
    return {"docs": d, "chunks": c}


def list_docs(limit=30, tag=None):
    init_db()
    conn = _connect()
    try:
        if tag:
            rows = conn.execute(
                """SELECT * FROM docs WHERE archived=0 AND tags LIKE ?
                   ORDER BY updated_at DESC, id DESC LIMIT ?""",
                (f'%"{str(tag).lower()}"%', limit),
            ).fetchall()
        else:
            rows = conn.execute(
                """SELECT * FROM docs WHERE archived=0
                   ORDER BY updated_at DESC, id DESC LIMIT ?""",
                (limit,),
            ).fetchall()
    finally:
        conn.close()
    return [
        {
            "id": r["id"], "uid": r["uid"], "title": r["title"],
            "tags": json.loads(r["tags"] or "[]"), "source": r["source"],
            "origin": r["origin"], "summary": r["summary"], "n_chunks": r["n_chunks"],
            "created_at": r["created_at"], "updated_at": r["updated_at"],
        }
        for r in rows
    ]


# ======================================================================
# CLI — tez tekshiruv
# ======================================================================

if __name__ == "__main__":
    import sys

    args = sys.argv[1:]
    if args and args[0] == "--stats":
        init_db()
        print(json.dumps(stats(), ensure_ascii=False, indent=2))
    elif len(args) >= 2 and args[0] == "--ingest":
        p = Path(args[1])
        print(ingest_file(p, source="cli", origin=p.name))
    elif len(args) >= 2 and args[0] == "--search":
        for r in search(" ".join(args[1:])):
            print(f"{r['title']} › {r['heading']} | {r['text'][:120]}")
    else:
        init_db()
        print("kb: init ok", json.dumps(stats(), ensure_ascii=False))
